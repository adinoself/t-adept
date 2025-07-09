import requests
import streamlit as st
import transcribe
import time
import sys
from zipfile import ZipFile
from time import sleep
import os
import pickle
from pathlib import Path
import streamlit_authenticator as stauth
from docx import Document
import base64

auth_key = st.secrets['auth_key']

st.header("Songa Prime Transcription")
hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)

# --- USER AUTHENTICATION ---
names = ["Joseph Modi", "Emmah Wavinya", "Fuji Cheruiyot", "Eva Kimani", "Yusuf Kariuki", "George Orembo"]
usernames = ["adinoself", "missdivine", "fujicheruiyot", "evakimani", "yusufkariuki", "georgeorembo"]
credentials = {"usernames": {}}

# LOAD HASHED PASSWORDS
file_path = Path(__file__).parent / "hidden_pw.pkl"
with file_path.open("rb") as file:
    hashed_passwords = pickle.load(file)

authenticator = stauth.Authenticate(names, usernames, hashed_passwords, "transcription", "abcdef", cookie_expiry_days=30)

name, authentication_status, username = authenticator.login("Login", "main")

if authentication_status == False:
    st.error("Username/password is incorrect")

if authentication_status == None:
    st.warning("Please enter your username and password")

if authentication_status:

    # ----SIDEBAR ---
    authenticator.logout("Logout", "sidebar")
    st.sidebar.title(f"Welcome, {name}")

    # 1 Upload file to AssemblyAI
    def get_url(auth_key, data):
        headers = {'authorization': auth_key}
        response = requests.post('https://api.assemblyai.com/v2/upload',
                                 headers=headers,
                                 data=data)
        url = response.json()["upload_url"]
        print("Uploaded File and got temporary URL to file")
        return url

    # 2 Request transcription with summary
    def get_transcribe_id(auth_key, url):
        endpoint = "https://api.assemblyai.com/v2/transcript"
        json = {
            "audio_url": url,
            "speaker_labels": True,
            "summarization": True,
            "summary_model": "conversational",
            "summary_type": "bullets",  # Enable summarization
            "auto_chapters": True,  # Required for summary to work!
        }
        headers = {
            "authorization": auth_key,
            "content-type": "application/json"
        }
        response = requests.post(endpoint, json=json, headers=headers)
        print("Made request and file is currently queued")
        print(response.json())
        return response.json()['id']

    # 3 Upload and start transcription
    def upload_file(fileObj):
        auth_key = st.secrets['auth_key']
        url = get_url(auth_key, fileObj)
        t_id = get_transcribe_id(auth_key, url)
        return auth_key, t_id

    # 4 Poll result
    def get_text(auth_key, transcribe_id):
        endpoint = f"https://api.assemblyai.com/v2/transcript/{transcribe_id}"
        headers = {"authorization": auth_key}
        result = requests.get(endpoint, headers=headers).json()
        return result

    # UI
    fileObject = st.file_uploader(label="Please upload your file")
    if fileObject:
        auth_key, transcribe_id = upload_file(fileObject)
        result = {}
        sleep_duration = 1
        percent_complete = 0
        progress_bar = st.progress(percent_complete)
        st.text("Currently in queue")

        while result.get("status") != "processing":
            percent_complete += sleep_duration
            time.sleep(sleep_duration)
            progress_bar.progress(percent_complete / 10)
            result = get_text(auth_key, transcribe_id)

        sleep_duration = 0.01
        for percent in range(percent_complete, 101):
            time.sleep(sleep_duration)
            progress_bar.progress(percent)

        with st.spinner("Processing....."):
            while result.get("status") != 'completed':
                result = get_text(auth_key, transcribe_id)

        # Display transcription
        st.success('Transcription Successful!')
        st.subheader("Transcribed Text")
        st.success(result['text'])

        # Write plain text transcript
        doc = Document()
        doc.add_paragraph(result["text"])
        doc.save('plain text transcript.docx')

        # With speaker labels
        doc = Document()
        for utterance in result["utterances"]:
            speaker = utterance["speaker"]
            text = utterance["text"]
            doc.add_paragraph(f"Speaker {speaker}: {text}")
        doc.save('transcript with speaker labels.docx')

        # Write SRT
        endpoint = f"https://api.assemblyai.com/v2/transcript/{transcribe_id}"
        headers = {"authorization": auth_key}
        srt_endpoint = endpoint + "/srt"
        srt_response = requests.get(srt_endpoint, headers=headers)
        with open("transcript with time stamps.txt", "w") as _file:
            _file.write(srt_response.text)

        # Write summary
        summary_text = result.get("summary", "Summary not available.")
        with open("summary.txt", "w") as f:
            f.write(str(summary_text))

        # Zip
        with ZipFile('transcription.zip', 'w') as zip_file:
            zip_file.write('plain text transcript.docx')
            zip_file.write('transcript with speaker labels.docx')
            zip_file.write('transcript with time stamps.txt')
            zip_file.write('summary.txt')

        # Prepare download links
        def file_download_link(filename, display_name, mime_type="application/octet-stream"):
            with open(filename, "rb") as file:
                encoded = base64.b64encode(file.read()).decode()
                return f'<a href="data:{mime_type};base64,{encoded}" download="{filename}">{display_name}</a>'

        st.subheader("Download Transcripts")
        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown(file_download_link("plain text transcript.docx", "Download Plain Text Transcript.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"), unsafe_allow_html=True)
        with col2:
            st.markdown(file_download_link("transcript with speaker labels.docx", "Download Transcript with Speaker Labels.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"), unsafe_allow_html=True)
        with col3:
            st.markdown(file_download_link("transcript with time stamps.txt", "Download Transcript with Time Stamps.txt", "text/plain"), unsafe_allow_html=True)

        # Summary
        st.subheader("Download Summary")
        st.markdown(file_download_link("summary.txt", "Download Summary.txt", "text/plain"), unsafe_allow_html=True)

        # All-in-one
        st.subheader("Download All Transcripts")
        st.markdown(file_download_link("transcription.zip", "Download All Transcripts as ZIP", "application/zip"), unsafe_allow_html=True)
